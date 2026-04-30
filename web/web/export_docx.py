import io
import os
import glob
import logging
import json
from datetime import datetime

from django.http import HttpResponse, Http404
from django.contrib.auth.models import User
from analysis.views import classify_malware_ml, _static_classify

from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

log = logging.getLogger(__name__)

# Colour Palette 
C_DARK    = RGBColor(0x0A, 0x16, 0x28)
C_LBLUE   = RGBColor(0x1D, 0x9B, 0xF0)
C_BODY    = RGBColor(0x1A, 0x23, 0x32)

C_TABLE_HDR_BG = "D4E8F5"  
C_KEY_BG       = "E8F0F8"  
C_ROW_EVEN     = "F4F8FC"
C_ROW_ODD      = "FFFFFF"

C_SEC_BG       = "EAF3FB"  
C_SEC_BAR      = "1D6FA4"  

C_DANGER  = RGBColor(0xC0, 0x39, 0x2B)
C_WARN    = RGBColor(0xB7, 0x77, 0x0D)
C_SAFE    = RGBColor(0x1A, 0x7A, 0x46)
C_FOOTER  = RGBColor(0x4A, 0x7A, 0x9A)
C_BLUE_TXT= RGBColor(0x25, 0x63, 0xEB)

CONTENT_W = 9638


# XML Low-level 
def _tbl_from_xml(container, rows: int, cols: int, width_cm: float = 17.0):
    from docx.table import Table
    from docx.oxml.table import CT_Tbl
    tbl_elem = CT_Tbl.new_tbl(rows, cols, Cm(width_cm))
    container._element.append(tbl_elem)
    tbl = Table(tbl_elem, container)
    tbl.style = 'Normal Table'
    return tbl

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def _set_cell_border(cell, **sides):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = OxmlElement("w:tcBorders")
        tcPr.append(b)
    else:
        b.clear()

    for side, (style, size, color) in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")  
        el.set(qn("w:color"), color.lstrip("#"))
        b.append(el)

def _no_border(cell):
    _set_cell_border(
        cell,
        top=("none", 0, "FFFFFF"), bottom=("none", 0, "FFFFFF"),
        left=("none", 0, "FFFFFF"), right=("none", 0, "FFFFFF"),
    )

def _thin_border(cell, color="A0C0D0"): 
    _set_cell_border(
        cell,
        top=("single", 4, color), bottom=("single", 4, color),
        left=("single", 4, color), right=("single", 4, color),
    )

def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)

def _set_col_widths(table, widths: list):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    idx = list(tbl).index(tblPr)
    tbl.insert(idx + 1, tblGrid)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            w = widths[i] if i < len(widths) else widths[-1]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)

def _set_row_cant_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), 'true')
    trPr.append(cantSplit)

def _para_bottom_border(para, color="1D6FA4", size=8):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bot)
    pPr.append(pBdr)

def _para_left_bar(para, color, size=24):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(left)
    pPr.append(pBdr)

def _para_bg(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    pPr.append(shd)

def _add_page_number(para):
    def _field(instr):
        r1 = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), "begin")
        r1.append(fc)
        para._p.append(r1)

        r2 = OxmlElement("w:r")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = instr
        r2.append(it)
        para._p.append(r2)

        r3 = OxmlElement("w:r")
        fe = OxmlElement("w:fldChar")
        fe.set(qn("w:fldCharType"), "end")
        r3.append(fe)
        para._p.append(r3)

    _field(" PAGE ")
    r_sep = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = " / "
    r_sep.append(t)
    para._p.append(r_sep)
    _field(" NUMPAGES ")

def _run(para, text, bold=False, italic=False, size_pt=10, color=None, font="Calibri"):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size_pt)
    if color:
        r.font.color.rgb = color
    return r

def _spacer(doc, pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pt)


# Header / Footer 
def _build_header(doc, analysis: dict, logo_path, now: str):
    sec = doc.sections[0]
    sec.different_first_page_header_footer = False
    hdr = sec.header
    hdr.is_linked_to_previous = False

    for p in list(hdr.paragraphs):
        p._element.getparent().remove(p._element)

    COL_W = [1500, 5500, 2638]
    tbl = _tbl_from_xml(hdr, 1, 3, 17.0)
    _set_col_widths(tbl, COL_W)
    for cell in tbl.rows[0].cells:
        _no_border(cell)

    lc = tbl.rows[0].cells[0]
    _set_cell_bg(lc, "0A1628")
    _set_cell_margins(lc, top=60, bottom=60, left=80, right=80)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = lc.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path and os.path.isfile(logo_path):
        try:
            lp.add_run().add_picture(logo_path, width=Cm(2.0))
        except Exception as exc:
            log.warning("logo load failed: %s", exc)
            _run(lp, "\u25c6A", bold=True, size_pt=20, color=C_LBLUE, font="Courier New")
    else:
        _run(lp, "\u25c6A", bold=True, size_pt=20, color=C_LBLUE, font="Courier New")

    tc2 = tbl.rows[0].cells[1]
    _set_cell_bg(tc2, "0A1628")
    _set_cell_margins(tc2, top=80, bottom=80, left=160, right=80)
    tc2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tp = tc2.paragraphs[0]
    tp.paragraph_format.space_after = Pt(2)
    _run(tp, "AMAL ", bold=True, size_pt=18, color=RGBColor(0xE8, 0xF4, 0xFF))
    _run(tp, "v4",   bold=True, size_pt=18, color=C_LBLUE)
    sub = tc2.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after  = Pt(0)
    _run(sub, "Automatic Malware Analyzer", size_pt=8, color=RGBColor(0x64, 0xA9, 0xD4))

    mc = tbl.rows[0].cells[2]
    _set_cell_bg(mc, "0D1E35")
    _set_cell_margins(mc, top=60, bottom=60, left=80, right=100)
    mc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _meta(label, value):
        mp = mc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        mp.paragraph_format.space_before = Pt(1)
        mp.paragraph_format.space_after  = Pt(1)
        _run(mp, label, bold=True, size_pt=8, color=RGBColor(0xC8, 0xDF, 0xF0))
        _run(mp, value, size_pt=8, color=RGBColor(0x94, 0xB8, 0xD0))

    info    = analysis.get("info", {})
    task_id = info.get("id", "?")
    _meta("Report Date:  ", now)
    _meta("Task ID:  ", f"{task_id}")
    fname = (analysis.get("target", {}).get("file") or {}).get("name", "")
    url   = (analysis.get("target") or {}).get("url", "")
    if fname:
        _meta("Sample:  ", fname[:28])
    elif url:
        _meta("URL:  ", url[:28])

    docno_p = mc.add_paragraph()
    docno_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    docno_p.paragraph_format.space_before = Pt(4)
    docno_p.paragraph_format.space_after  = Pt(0)
    _run(docno_p, f"DOC-AMAL-{task_id}", bold=True, size_pt=9, color=C_LBLUE)

def _build_footer(doc, analysis: dict, now: str):
    sec = doc.sections[0]
    ftr = sec.footer
    ftr.is_linked_to_previous = False

    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)

    COL_W = [3300, 3038, 3300]
    tbl = _tbl_from_xml(ftr, 1, 3, 17.0)
    _set_col_widths(tbl, COL_W)
    for cell in tbl.rows[0].cells:
        _no_border(cell)
        _set_cell_margins(cell, top=30, bottom=30, left=60, right=60)

    lp = tbl.rows[0].cells[0].paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(lp, "AMAL v4 \u2014 Automatic Malware Analyzer", size_pt=7.5, color=C_FOOTER)

    cp = tbl.rows[0].cells[1].paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(cp, "Badan Siber dan Sandi Negara", bold=True, size_pt=8, color=C_LBLUE)

    task_id = analysis.get("info", {}).get("id", "?")
    rp = tbl.rows[0].cells[2].paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(rp, f"Task {task_id} \u2014 Page ", size_pt=7.5, color=C_FOOTER)
    _add_page_number(rp)


# Heading helpers 
def _section_heading(doc, title: str):
    count = getattr(doc, '_amal_heading_count', 0)
    if count > 0 and count % 2 == 0:
        doc.add_page_break()
    setattr(doc, '_amal_heading_count', count + 1)
    
    p = doc.add_heading("", level=1)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(14)
    p.paragraph_format.keep_with_next = True 
    
    _para_left_bar(p, C_SEC_BAR, 24)
    _para_bg(p, C_SEC_BG)
    _run(p, f" \u25a0  {title}", bold=True, size_pt=12, color=C_DARK)
    
    return p

def _sub_heading(doc, title: str):
    p = doc.add_heading("", level=2)
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(10)
    p.paragraph_format.keep_with_next = True 
    
    _para_bottom_border(p, "B8D4E8", 4)
    _run(p, title, bold=True, size_pt=10.5, color=RGBColor(0x1A, 0x3A, 0x5C))
    return p

# Table helpers 
def _kv_table(doc, rows: list, key_w=2800, val_w=6838):
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER 
    _set_col_widths(t, [key_w, val_w])
    
    for key, val in rows:
        row = t.add_row()
        _set_row_cant_split(row) 
        
        kc = row.cells[0]
        _set_cell_bg(kc, C_KEY_BG); _set_cell_margins(kc); _thin_border(kc)
        p_key = kc.paragraphs[0]
        p_key.paragraph_format.keep_with_next = True
        _run(p_key, str(key), bold=True, size_pt=9, color=C_DARK)
        
        vc = row.cells[1]
        _set_cell_margins(vc); _thin_border(vc)
        _run(vc.paragraphs[0], str(val) if val is not None else "\u2014", size_pt=9, color=C_BODY)
    
    _spacer(doc, pt=12) 
    return t

def _data_table(doc, headers: list, rows_data: list, col_widths: list = None, mono_cols: list = None):
    ncols = len(headers)
    if not col_widths:
        w = CONTENT_W // ncols
        col_widths = [w] * ncols
    t = doc.add_table(rows=0, cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(t, col_widths)
    mono_cols = mono_cols or []

    hr = t.add_row()
    _set_row_cant_split(hr)
    for i, hdr in enumerate(headers):
        cell = hr.cells[i]
        _set_cell_bg(cell, C_TABLE_HDR_BG); _set_cell_margins(cell)
        _thin_border(cell)
        
        p = cell.paragraphs[0]
        p.paragraph_format.keep_with_next = True
        _run(p, hdr, bold=True, size_pt=9, color=C_DARK)

    for ri, row_data in enumerate(rows_data):
        row = t.add_row()
        _set_row_cant_split(row)
        bg = C_ROW_EVEN if ri % 2 == 0 else C_ROW_ODD
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg); _set_cell_margins(cell); _thin_border(cell)
            _run(cell.paragraphs[0], str(val) if val is not None else "\u2014",
                 size_pt=8.5 if ci in mono_cols else 9, color=C_BODY,
                 font="Consolas" if ci in mono_cols else "Calibri")
            
    _spacer(doc, pt=12)
    return t


# Data helpers 
def _resolve_custom(analysis: dict):
    info = analysis.get("info", {})
    raw = info.get("custom", {})

    if isinstance(raw, str) and raw.strip():
        try:
            raw = json.loads(raw)
        except Exception:
            raw = {"submitter": raw}
    elif not isinstance(raw, dict):
        raw = {}

    submitter = raw.get("submitter", "")
    if not submitter or submitter == "anonymous":
        submitter = (
            str(info.get("owner", "") or "").strip()
            or str(info.get("user_id", "") or "").strip()
        )
        if submitter:
            raw["submitter"] = submitter

    if not submitter or submitter == "anonymous":
        return None

    from django.contrib.auth.models import User
    try:
        if str(submitter).isdigit():
            user = User.objects.filter(id=int(submitter)).first()
        else:
            user = User.objects.filter(username=submitter).first()

        if user:
            raw["submitter_name"] = user.get_full_name() or user.username
            raw["submitter"]      = user.username
            raw["email"]          = user.email
            if hasattr(user, "userprofile"):
                profile = user.userprofile
                raw["organization"] = getattr(profile, "organization", "")
                raw["unit"]         = getattr(profile, "unit", "")
    except Exception as exc:
        log.warning("_resolve_custom user lookup failed: %s", exc)

    return {
        "full_name":    raw.get("submitter_name") or raw.get("submitter") or "Anonymous",
        "submitter":    raw.get("submitter", "anonymous"),
        "email":        raw.get("email", ""),
        "organization": raw.get("organization", ""),
        "unit":         raw.get("unit", ""),
    }

    
def _resolve_screenshots(analysis: dict) -> list:
    from django.conf import settings
    task_id  = analysis.get("info", {}).get("id")
    shots_dir = os.path.join(
        getattr(settings, "CUCKOO_PATH", "/opt/CAPEv2"),
        "storage", "analyses", str(task_id), "shots",
    )
    if not os.path.isdir(shots_dir): return []
    paths = []
    for pat in ("*.jpg", "*.jpeg", "*.png"):
        paths.extend(sorted(glob.glob(os.path.join(shots_dir, pat))))
    return sorted(paths)[-10:]

def _resolve_logo_path():
    paths = [
        "/opt/CAPEv2/web/static/img/BSSN.png",
        "/opt/CAPEv2/data/html/graphic/logo.png",
        "/opt/CAPEv2/data/html/graphic/logo.jpg",
    ]
    for p in paths:
        if os.path.isfile(p): return p
    return None

def _resolve_dns_answers(analysis: dict):
    for entry in analysis.get("network", {}).get("dns", []):
        raw = entry.get("answers", [])
        parts = []
        for a in raw:
            if isinstance(a, dict): 
                rtype = a.get("type", "")
                rdata = a.get("data") or a.get("address") or str(next(iter(a.values()), "?"))
                parts.append(f"{rtype} {rdata}" if rtype else rdata)
            else: 
                parts.append(str(a))
        entry["answers_str"] = "\n".join(parts) if parts else "\u2014"

def _resolve_radar(analysis: dict):
    return analysis.get("radar_labels", []), analysis.get("radar_data", [])

def _resolve_radar_combined(labels, data) -> list:
    combined = sorted(zip(labels, data), key=lambda x: x[1], reverse=True)
    result = []
    for label, score in combined:
        bl = int((score / 100) * 25)
        result.append({"label": label, "score": score, "bar": "\u2588" * bl + "\u2591" * (25 - bl)})
    return result

def _resolve_stats_total(analysis: dict) -> dict:
    stats = analysis.get("statistics", {})
    def _s(key): return sum(float(x.get("time", 0)) for x in stats.get(key, []) if x.get("time"))
    return {k: round(_s(k), 2) for k in ("processing", "signatures", "reporting")}

def _resolve_behavior_categories(analysis: dict) -> list:
    summary = analysis.get("behavior", {}).get("summary", {})
    return [
        ("Accessed Files",     summary.get("files", [])),
        ("Executed Commands",  summary.get("executed_commands", [])),
        ("Modified Files",     summary.get("write_files", [])),
        ("Deleted Files",      summary.get("delete_files", [])),
        ("Read Files",         summary.get("read_files", [])),
        ("Modified Registry",  summary.get("write_keys", [])),
        ("Deleted Registry",   summary.get("delete_keys", [])),
        ("Mutexes",            summary.get("mutexes", [])),
        ("Created Services",   summary.get("created_services", [])),
        ("Started Services",   summary.get("started_services", [])),
        ("Resolved APIs",      summary.get("resolved_apis", [])),
    ]

def _inject_entropy_bar(analysis: dict):
    bh = analysis.get("bytehist_analysis", {})
    if not bh: return
    pct = bh.get("entropy_pct", 0)
    bl  = int(pct / 100 * 30)
    bh["entropy_bar"] = "\u2588" * bl + "\u2591" * (30 - bl)

def _resolve_capa(analysis: dict) -> list:
    capa = analysis.get("capa_summary", {})
    if not capa: return []
    result = []
    for sk in ("CAPABILITY", "ATTCK", "MBC"):
        sd = capa.get(sk, {})
        if not sd: continue
        items = []
        for ns, content in sd.items():
            if isinstance(content, list): entries = [str(x) for x in content]
            elif isinstance(content, dict): entries = [f"{k}: {v}" for k, v in content.items()]
            else: entries = [str(content)]
            items.append({"ns": ns, "entries": entries})
        if items: result.append({"section": sk, "items": items})
    return result

def _score_color(score) -> RGBColor:
    try: s = float(score)
    except: return C_BODY
    return C_DANGER if s > 6 else (C_WARN if s > 2 else C_SAFE)

def _bar_color(score: int) -> RGBColor:
    return C_DANGER if score > 60 else (C_WARN if score > 30 else C_SAFE)

def _sig_label(severity: int):
    if severity >= 3: return "CRITICAL", C_DANGER
    if severity == 2: return "WARN",     C_WARN
    if severity == 1: return "INFO",     C_BLUE_TXT
    return "LOW", C_SAFE


# Section Builders 
def _add_analysis_details(doc, analysis: dict):
    _section_heading(doc, "Analysis Details")

    detections = analysis.get("detections")
    if detections:
        if isinstance(detections, str) and detections.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _run(p, "Detections:   ", bold=True, size_pt=10)
            _run(p, detections, bold=True, size_pt=10, color=C_DANGER)

        elif isinstance(detections, list) and detections:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _run(p, "Detections:   ", bold=True, size_pt=10)
            for d in detections:
                family = d.get("family", str(d)) if isinstance(d, dict) else str(d)
                pd = doc.add_paragraph()
                pd.paragraph_format.left_indent = Cm(0.8)
                pd.paragraph_format.space_after = Pt(2)
                _run(pd, f"● {family}", bold=True, size_pt=9, color=C_DANGER)
            _spacer(doc, pt=4)

    # MalScore 
    ms = analysis.get("malscore")
    if ms is not None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, "MalScore:   ", bold=True, size_pt=10)
        _run(p, f"{ms} / 10", bold=True, size_pt=20, color=_score_color(ms))

    # MalStatus 
    malstatus = analysis.get("malstatus")
    if malstatus:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, "Status:   ", bold=True, size_pt=10)
        col = (C_DANGER if malstatus == "Malicious" else C_WARN if malstatus == "Suspicious" else C_SAFE)
        _run(p, f"[ {malstatus} ]", bold=True, size_pt=10, color=col)

    # Info Table 
    info = analysis.get("info", {})
    _kv_table(doc, [
        ("Category",  (info.get("category") or "\u2014").upper()),
        ("Package",   info.get("package") or "\u2014"),
        ("Started",   info.get("started") or "\u2014"),
        ("Completed", info.get("ended") or "\u2014"),
        ("Duration",  f"{info.get('duration', '\u2014')}s"),
        ("Route",     info.get("route") or "\u2014"),
    ])

def _add_uploader_info(doc, analysis: dict):
    custom_info = _resolve_custom(analysis)
    if not custom_info:
        return

    _section_heading(doc, "Uploader Information")

    from django.contrib.auth.models import User
    submitter_uname = custom_info.get("submitter", "")
    status = "Unknown"

    if submitter_uname and submitter_uname != "anonymous":
        user = User.objects.filter(username=submitter_uname).first()
        if user:
            status_parts = ["Active" if user.is_active else "Inactive"]
            if user.is_superuser:
                status_parts.append("Superuser")
            elif user.is_staff:
                status_parts.append("Staff")
            status = " | ".join(status_parts)

    rows = [
        ("Full Name",       custom_info.get("full_name", "\u2014")),
        ("Username",        f"@{custom_info.get('submitter', '\u2014')}"),
        ("Email",           custom_info.get("email") or "\u2014"),
        ("Organization",    custom_info.get("organization") or "\u2014"),
        ("Unit / Division", custom_info.get("unit") or "\u2014"),
        ("Status",          status),
    ]
    _kv_table(doc, rows)

def _add_machine_info(doc, analysis: dict):
    machine = (analysis.get("info") or {}).get("machine") or {}
    if not machine.get("name"): return
    _section_heading(doc, "Machine Information")
    rows = [
        ("Name",        machine.get("name", "\u2014")),
        ("Label",       machine.get("label", "\u2014")),
        ("Manager",     machine.get("manager", "\u2014")),
    ]
    _kv_table(doc, rows)

def _add_malware_classification(doc, analysis: dict):
    report_category = analysis.get("info", {}).get("category", "file")

    if report_category == "static":
        radar_labels, radar_data, _, _, _ = _static_classify(analysis)
    elif report_category == "pcap":
        radar_labels, radar_data = [], []
    else:
        radar_labels, radar_data, _, _, _ = classify_malware_ml(analysis)

    if not radar_labels or not radar_data:
        return

    max_score = max(radar_data) if radar_data else 1
    radar_combined = []
    for lbl, score in zip(radar_labels, radar_data):
        bl = int((score / max_score) * 25) if max_score > 0 else 0
        bar_visual = "\u2588" * bl + "\u2591" * (25 - bl)
        radar_combined.append({"label": lbl, "score": score, "bar": bar_visual})

    _section_heading(doc, "Malware Classification")
    t = doc.add_table(rows=0, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(t, [2800, 900, 5938])

    hr = t.add_row()
    _set_row_cant_split(hr)
    for i, hdr in enumerate(["Category", "Score", "Bar"]):
        cell = hr.cells[i]
        _set_cell_bg(cell, C_TABLE_HDR_BG)
        _set_cell_margins(cell)
        _thin_border(cell)
        _run(cell.paragraphs[0], hdr, bold=True, size_pt=9, color=C_DARK)

    for ri, item in enumerate(radar_combined):
        row = t.add_row()
        _set_row_cant_split(row)
        bg = C_ROW_EVEN if ri % 2 == 0 else C_ROW_ODD
        score_val = item["score"]

        vals = [
            (item["label"], False, True,  False),
            (f"{score_val:.1f}", False, False, False),
            (item["bar"],   True,  False, True),
        ]
        for ci, (val, mono, bold, is_bar) in enumerate(vals):
            cell = row.cells[ci]
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell)
            _thin_border(cell)
            clr = _bar_color(score_val) if is_bar else C_BODY
            _run(cell.paragraphs[0], val,
                 bold=bold,
                 size_pt=8.5 if mono else 9,
                 color=clr,
                 font="Consolas" if mono else "Calibri")
    _spacer(doc, pt=12)

def _add_malware_config(doc, analysis: dict):
    configs = analysis.get("malware_conf")
    if not configs: return
    _section_heading(doc, "Malware Configuration")
    for cb in configs:
        for family, config in cb.items():
            if family.startswith("_"): continue
            _sub_heading(doc, family)
            rows = []
            for key, val in config.items():
                if key == "raw":
                    for item in val:
                        for sk, sv in item.items(): rows.append((sk, str(sv)))
                else: rows.append((key, str(val)))
            _kv_table(doc, rows)

def _add_file_details(doc, analysis: dict):
    f = (analysis.get("target") or {}).get("file") or {}
    if not f: return
    _section_heading(doc, "File Details")
    rows = []
    for label, key in [
        ("File Name", "name"), ("Size", "size"), ("Type", "type"),
        ("MD5", "md5"), ("SHA1", "sha1"), ("SHA256", "sha256"),
        ("CRC32", "crc32"), ("MIME", "mime"),
    ]:
        v = f.get(key)
        if v: rows.append((label, f"{v} bytes" if key == "size" else str(v)))
    
    _kv_table(doc, rows, key_w=2000, val_w=7638)

def _add_capa(doc, capa_resolved: list):
    if not capa_resolved: return
    _section_heading(doc, "CAPA Analysis Summary")
    for block in capa_resolved:
        _sub_heading(doc, block["section"])
        rows_data = [[item["ns"], "\n".join(item["entries"])] for item in block["items"]]
        _data_table(doc, ["Namespace", "Items"], rows_data, col_widths=[3000, 6638])

def _add_curtain(doc, analysis: dict):
    curtain = analysis.get("curtain")
    if not curtain: return
    _section_heading(doc, "Curtain (PowerShell Events)")
    for pid, pdata in curtain.items():
        behaviors = ", ".join(pdata.get("behaviors", []))
        _sub_heading(doc, f"PID: {pid}  {behaviors}")
        rows_data = []
        for entry in pdata.get("events", []):
            for event, msg in entry.items():
                orig = msg.get("original", str(msg)) if isinstance(msg, dict) else str(msg)
                rows_data.append([event, orig[:500]])
        _data_table(doc, ["Event", "Original"], rows_data, col_widths=[2600, 7038])

def _add_mitre(doc, analysis: dict):
    mitre = analysis.get("mitre_attck")
    if not mitre: return
    _section_heading(doc, "MITRE ATT&CK")
    for tactic, ttps in mitre.items():
        _sub_heading(doc, tactic)
        rows_data = [[ttp.get("t_id", ""), ttp.get("ttp_name", ""), "\n".join(ttp.get("signature", []))] for ttp in ttps]
        _data_table(doc, ["ID", "Technique", "Signatures"], rows_data, col_widths=[1200, 2800, 5638])

def _add_statistics(doc, analysis: dict, stats_total: dict):
    stats = analysis.get("statistics")
    if not stats: return
    _section_heading(doc, "Statistics")
    for key, label in [("processing", "Processing"), ("signatures", "Signatures"), ("reporting",  "Reporting")]:
        items = stats.get(key, [])
        if not items: continue
        _sub_heading(doc, f"{label} \u2014 Total: {stats_total.get(key, 0)}s")
        sorted_items = sorted(items, key=lambda x: float(x.get("time", 0) or 0), reverse=True)
        rows_data = [[r.get("name", ""), str(r.get("time", ""))] for r in sorted_items if r.get("time")]
        _data_table(doc, ["Module", "Time (s)"], rows_data, col_widths=[7000, 2638])

def _add_bytehist(doc, analysis: dict):
    bh = analysis.get("bytehist_analysis", {})
    if not bh: return
    _section_heading(doc, "Byte Statistics & Entropy Analysis")

    _kv_table(doc, [
        ("Shannon Entropy",  f"{bh.get('entropy', 0):.4f} bits/byte"),
        ("χ² Statistic",     f"{bh.get('chi_square', 0):.1f}  df=255  {str(bh.get('chi_verdict', '')).upper()}"),
        ("NULL Bytes",       f"{bh.get('null_pct', 0):.2f}% of total bytes"),
        ("Printable ASCII",  f"{bh.get('printable_pct', 0):.2f}% of total bytes"),
        ("File Size",        f"{bh.get('file_size', 0) / (1024*1024):.1f} MB  —  {bh.get('file_size', 0):,} bytes"),
    ])

    cls      = bh.get("classification", "")
    cls_desc = bh.get("classification_desc", "")
    cls_conf = bh.get("confidence", 0)
    if cls:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, cls, bold=True, size_pt=11, color=C_DARK)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        _run(p2, cls_desc, size_pt=9, color=C_BODY)

        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(8)
        _run(p3, "Confidence ", bold=True, size_pt=9, color=C_DARK)
        _run(p3, f"{cls_conf}%", size_pt=9, color=_bar_color(cls_conf))

    indicators = bh.get("indicators", [])
    if indicators:
        level_color = {"high": C_DANGER, "medium": C_WARN, "low": C_SAFE, "info": C_BLUE_TXT}
        for ind in indicators:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(3)
            lvl = ind.get("level", "info")
            _run(p, "● ", bold=True, size_pt=9, color=level_color.get(lvl, C_BODY))
            _run(p, ind.get("text", ""), size_pt=9, color=C_BODY)
        _spacer(doc, pt=8)

    regions = bh.get("regions", {})
    if regions:
        _sub_heading(doc, "Byte Region Breakdown")
        _data_table(doc, ["Region", "%"], [
            ["NULL",    f"{regions.get('null_pct',    0):.1f}%"],
            ["Control", f"{regions.get('control_pct', 0):.1f}%"],
            ["ASCII",   f"{regions.get('ascii_pct',   0):.1f}%"],
            ["DEL",     f"{regions.get('del_pct',     0):.1f}%"],
            ["High",    f"{regions.get('high_pct',    0):.1f}%"],
        ], col_widths=[4819, 4819])

    top_bytes = bh.get("top_bytes", [])
    if top_bytes:
        _sub_heading(doc, "Most Frequent Bytes")
        rows = []
        for b in top_bytes[:10]:
            byte_val = b.get("decimal", 0)
            rows.append([
                str(byte_val),
                b.get("hex", hex(byte_val)),
                f"{b.get('count', 0):,}",
                f"{b.get('pct', 0):.2f}%",
            ])
        _data_table(doc, ["Byte", "Hex", "Count", "%"], rows,
                    col_widths=[1500, 1500, 3000, 3638], mono_cols=[1])

    rare_bytes = bh.get("rare_bytes", [])
    if rare_bytes:
        _sub_heading(doc, "Rarest Bytes")
        rows = []
        for b in rare_bytes[:10]:
            byte_val = b.get("decimal", 0)
            rows.append([
                str(byte_val),
                b.get("hex", hex(byte_val)),
                f"{b.get('count', 0):,}",
                f"{b.get('pct', 0):.4f}%",
            ])
        _data_table(doc, ["Byte", "Hex", "Count", "%"], rows,
                    col_widths=[1500, 1500, 3000, 3638], mono_cols=[1])
def _add_signatures(doc, analysis: dict):
    sigs = analysis.get("signatures", [])
    if not sigs: return
    _section_heading(doc, f"Signatures ({len(sigs)})")
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(t, [1400, 8238])
    hr = t.add_row()
    _set_row_cant_split(hr)
    for i, hdr in enumerate(["Severity", "Description"]):
        cell = hr.cells[i]
        _set_cell_bg(cell, C_TABLE_HDR_BG); _set_cell_margins(cell); _thin_border(cell)
        _run(cell.paragraphs[0], hdr, bold=True, size_pt=9, color=C_DARK)

    SEV_HEX = {"C_DANGER": "C0392B", "C_WARN": "E67E22", "C_BLUE": "2980B9", "C_SAFE": "27AE60"}
    SEV_KEY = {C_DANGER: "C_DANGER", C_WARN: "C_WARN", C_BLUE_TXT: "C_BLUE", C_SAFE: "C_SAFE"}
    for ri, sig in enumerate(sigs):
        row = t.add_row()
        _set_row_cant_split(row)
        bg = C_ROW_EVEN if ri % 2 == 0 else C_ROW_ODD
        label, color = _sig_label(sig.get("severity", 0))
        hex_key = SEV_HEX.get(SEV_KEY.get(color, "C_BLUE"), "2980B9")

        sc = row.cells[0]
        _set_cell_bg(sc, bg); _set_cell_margins(sc)
        _set_cell_border(sc, left=("single", 12, hex_key), top=("single", 4, "D0DDE8"), bottom=("single", 4, "D0DDE8"), right=("single", 4, "D0DDE8"))
        _run(sc.paragraphs[0], label, bold=True, size_pt=8.5, color=color)

        dc = row.cells[1]
        _set_cell_bg(dc, bg); _set_cell_margins(dc); _thin_border(dc)
        _run(dc.paragraphs[0], sig.get("description", "\u2014"), size_pt=9, color=C_BODY)
    _spacer(doc, pt=12)

def _add_network(doc, analysis: dict):
    import re
    network = analysis.get("network", {})
    hosts = network.get("hosts", [])
    dns = network.get("dns", [])
    
    if not hosts and not dns: return
    _section_heading(doc, "Network Activity")

    ip_regex = re.compile(r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b")

    domains = network.get("domains", [])
    domain_to_ip = {
        d["domain"]: d["ip"]
        for d in domains
        if d.get("domain") and d.get("ip")
    }

    all_dns_ips = set()
    for d in dns:
        for ans in d.get("answers", []):
            val = str(ans.get("data", ans.get("address", ""))) if isinstance(ans, dict) else str(ans)
            for ip in ip_regex.findall(val):
                all_dns_ips.add(ip)

    if hosts:
        _sub_heading(doc, "Hosts")
        host_rows = []
        for h in hosts:
            ip = h.get("ip", "")
            country = h.get("country_name", "\u2014")
            asn = str(h.get("asn", "\u2014"))
            direct = "N" if ip in all_dns_ips else "Y"
            host_rows.append([direct, ip, country, asn])
        _data_table(doc, ["Direct", "IP", "Country Name", "ASN"],
                    host_rows, col_widths=[1200, 3000, 3000, 2438], mono_cols=[1])

    if dns:
        _sub_heading(doc, "DNS")
        dns_rows = []
        for d in dns:
            req_name = d.get("request", "")
            answers = d.get("answers_str", "\u2014")

            post_analysis = domain_to_ip.get(req_name, "\u2014")

            dns_rows.append([req_name, answers, post_analysis])

        _data_table(doc, ["Name", "Response", "Post-Analysis Lookup"],
                    dns_rows, col_widths=[2800, 4638, 2200], mono_cols=[0])

def _add_screenshots(doc, screenshots: list):
    if not screenshots: return
    _section_heading(doc, f"Screenshots ({len(screenshots)})")
    for i in range(0, len(screenshots), 2):
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_col_widths(t, [4819, 4819])
        for side in range(2):
            idx  = i + side
            cell = t.rows[0].cells[side]
            _set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            _thin_border(cell)
            _set_cell_bg(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if idx < len(screenshots):
                try:
                    p.add_run().add_picture(screenshots[idx], width=Inches(3.0))
                except Exception as exc:
                    log.warning("screenshot %s skip — %s", screenshots[idx], exc)
        _spacer(doc)

# Limit
def _add_behavior_summary(doc, behavior_categories: list):
    if not any(items for _, items in behavior_categories): return
    _section_heading(doc, "Summary")
    for label, items in behavior_categories:
        if not items: continue
        _sub_heading(doc, f"{label} \u2014 {len(items)} item(s)")
        
        limit = min(len(items), 20)
        for i, item in enumerate(items[:limit]):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            if i < 3: 
                p.paragraph_format.keep_with_next = True
            _run(p, str(item), size_pt=8.5, color=C_BODY, font="Consolas")
            
            if i == limit - 1 and len(items) <= 20:
                p.paragraph_format.space_after = Pt(12)
                _para_bottom_border(p, "D0DDE8", size=4) 

        if len(items) > 20:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(12)
            _run(p, f"\u2026 and {len(items) - 20} more items", italic=True, size_pt=8.5, color=C_FOOTER)
            _para_bottom_border(p, "D0DDE8", size=4)

# show all
"""def _add_behavior_summary(doc, behavior_categories: list):
    if not any(items for _, items in behavior_categories): return
    _section_heading(doc, "Summary")

    for label, items in behavior_categories:
        if not items: continue
        _sub_heading(doc, f"{label} \u2014 {len(items)} item(s)")

        for i, item in enumerate(items):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            if i < 3:
                p.paragraph_format.keep_with_next = True
            _run(p, str(item), size_pt=8.5, color=C_BODY, font="Consolas")

        # Bottom border setelah item terakhir
        p.paragraph_format.space_after = Pt(12)
        _para_bottom_border(p, "D0DDE8", size=4) """

def _add_watermark(doc, logo_path, opacity: float = 0.08):
    if not logo_path or not os.path.isfile(logo_path):
        return

    from lxml import etree
    from io import BytesIO
    from PIL import Image as PilImage
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    try:
        img = PilImage.open(logo_path).convert("RGBA")
        r, g, b, a = img.split()
        a = a.point(lambda x: int(x * opacity))   
        img = PilImage.merge("RGBA", (r, g, b, a))
        buf = BytesIO()
        img.save(buf, format="PNG")
        img_bytes = buf.getvalue()
        w_px, h_px = img.size
    except Exception as exc:
        log.warning("watermark image processing failed: %s", exc)
        return

    target_w_cm = 10.0
    aspect      = h_px / w_px if w_px > 0 else 1.0
    target_h_cm = target_w_cm * aspect
    w_pt        = target_w_cm * 28.35
    h_pt        = target_h_cm * 28.35

    section = doc.sections[0]
    header  = section.header

    img_part = Part(
        PackURI("/word/media/wm_logo.png"),
        "image/png",
        img_bytes,
        doc.part.package,
    )
    rId = header.part.relate_to(
        img_part,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image",
    )

    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    VNS = "urn:schemas-microsoft-com:vml"
    ONS = "urn:schemas-microsoft-com:office:office"
    RNS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    W10 = "urn:schemas-microsoft-com:office:word"

    wm_xml = (
        f'<w:p xmlns:w="{WNS}">'
          f'<w:pPr><w:jc w:val="center"/></w:pPr>'
          f'<w:r><w:rPr><w:noProof/></w:rPr>'
          f'<w:pict>'
            f'<v:shape xmlns:v="{VNS}" xmlns:o="{ONS}" xmlns:r="{RNS}" '
              f'id="watermark1" type="#_x0000_t75" o:allowincell="f" '
              f'style="position:absolute;margin-left:0;margin-top:0;'
                     f'width:{w_pt:.1f}pt;height:{h_pt:.1f}pt;'
                     f'z-index:-251654144;'
                     f'mso-position-horizontal:center;'
                     f'mso-position-horizontal-relative:margin;'
                     f'mso-position-vertical:center;'
                     f'mso-position-vertical-relative:margin">'
              f'<v:imagedata r:id="{rId}" o:title="watermark"/>'
              f'<w10:wrap xmlns:w10="{W10}" type="none"/>'
            f'</v:shape>'
          f'</w:pict>'
          f'</w:r>'
        f'</w:p>'
    )

    wm_element = etree.fromstring(wm_xml)
    header._element.append(wm_element)

# Main entry 
def build_analysis_docx(analysis: dict) -> bytes:
    _resolve_dns_answers(analysis)
    _inject_entropy_bar(analysis)
    labels, data  = _resolve_radar(analysis)
    radar_combined = _resolve_radar_combined(labels, data)
    screenshots   = _resolve_screenshots(analysis)
    stats_total   = _resolve_stats_total(analysis)
    behavior_cats = _resolve_behavior_categories(analysis)
    logo_path     = _resolve_logo_path()
    capa_resolved = _resolve_capa(analysis)
    now           = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    doc = Document()
    
    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
        sec.left_margin = sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(5.2)  
        sec.bottom_margin = Cm(2.5)
        sec.header_distance = Cm(1.2)
        sec.footer_distance = Cm(1.0)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    _build_header(doc, analysis, logo_path, now)
    _build_footer(doc, analysis, now)
    _add_watermark(doc, logo_path, opacity=0.10)

    # Call Sequence
    _add_analysis_details(doc, analysis)    
    _add_uploader_info(doc, analysis) 
    _add_malware_classification(doc, analysis)       
    _add_machine_info(doc, analysis)         
    _add_file_details(doc, analysis)          
    
    _add_malware_config(doc, analysis)
    _add_capa(doc, capa_resolved)
    _add_curtain(doc, analysis)
    _add_mitre(doc, analysis)
    _add_statistics(doc, analysis, stats_total)
    _add_bytehist(doc, analysis)
    _add_signatures(doc, analysis)
    _add_network(doc, analysis)
    _add_screenshots(doc, screenshots)
    
    _add_behavior_summary(doc, behavior_cats) 

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Data Loader 
def _load_analysis_for_docx(task_id: int) -> dict:
    from analysis.views import (
        mongo_find_one, enabledconf,
        split_signature_calls, _inject_bytehist_into_report,
    )
    EXCLUDE = {"dropped": 0, "CAPE.payloads": 0, "procdump": 0,
               "procmemory": 0, "behavior.processes": 0, "memory": 0}
    report = None

    if enabledconf.get("mongodb"):
        report = mongo_find_one("analysis", {"info.id": int(task_id)}, EXCLUDE, sort=[("_id", -1)])
    elif enabledconf.get("elasticsearchdb"):
        try:
            from analysis.views import es, get_analysis_index, get_query_by_info_id
            hits = es.search(
                index=get_analysis_index(),
                query=get_query_by_info_id(task_id),
                _source={"excludes": list(EXCLUDE.keys())},
            )["hits"]["hits"]
            report = hits[0]["_source"] if hits else None
        except Exception as exc:
            log.warning("export_docx: ES query failed — %s", exc)

    if not report:
        raise Http404(f"Analysis {task_id} not found")

    report = split_signature_calls(report)
    _inject_bytehist_into_report(report)
    _inject_radar_into_report(report)
    return report


def _inject_radar_into_report(report: dict):
    if report.get("radar_labels"):
        return
    CATEGORIES = {
        "Persistence":       ["persist", "autorun", "startup", "registry", "service", "boot"],
        "Injection":         ["inject", "hollow", "shellcode", "dll_inject"],
        "Evasion":           ["evasion", "anti", "sandbox", "vm", "debug", "obfuscat", "pack"],
        "Ransomware":        ["ransom", "encrypt", "shadow", "vssadmin", "wbadmin"],
        "Credential Access": ["credential", "password", "lsass", "mimikatz", "keylog", "token"],
        "Network":           ["network", "http", "dns", "connect", "download", "upload", "c2"],
        "Trojan":            ["trojan", "rat", "backdoor", "remote", "reverse"],
        "Info Stealer":      ["steal", "stealer", "exfil", "spyware"],
        "Dropper":           ["drop", "dropper", "downloader", "loader"],
        "Discovery":         ["discovery", "enum", "recon", "scan"],
    }
    scores = {cat: 0 for cat in CATEGORIES}
    for sig in report.get("signatures", []):
        weight = {0: 1, 1: 2, 2: 5, 3: 10}.get(sig.get("severity", 1), 2)
        name = (sig.get("name", "") + " " + sig.get("description", "") + " " + sig.get("category", "")).lower()
        for cat, kws in CATEGORIES.items():
            if any(kw in name for kw in kws):
                scores[cat] = min(scores[cat] + weight, 100)
    filtered = {k: v for k, v in scores.items() if v > 0}
    report["radar_labels"] = list(filtered.keys())
    report["radar_data"]   = list(filtered.values())


# Django View 
def export_analysis_docx(request, task_id: int):
    from django.conf import settings
    if getattr(settings, "WEB_AUTHENTICATION", False) and not request.user.is_authenticated:
        from django.contrib.auth.views import redirect_to_login
        return redirect_to_login(request.get_full_path())

    if not getattr(settings, "ALLOW_DL_REPORTS_TO_ALL", True):
        if request.user.is_anonymous:
            raise Http404("Login required")
        profile = getattr(request.user, "userprofile", None)
        if profile and not getattr(profile, "reports", True):
            raise Http404("No permission")

    try:
        analysis = _load_analysis_for_docx(task_id)
    except Http404:
        raise
    except Exception as exc:
        log.exception("export_docx: load failed task %s", task_id)
        return HttpResponse(f"Error loading analysis: {exc}", status=500)

    try:
        docx_bytes = build_analysis_docx(analysis)
    except Exception as exc:
        log.exception("export_docx: build failed task %s", task_id)
        return HttpResponse(f"Error generating DOCX: {exc}", status=500)

    fname = f"amal_analysis_{task_id}.docx"
    resp  = HttpResponse(
        docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{fname}"'
    resp["Content-Length"]      = len(docx_bytes)
    return resp